"""Generate a single consolidated Word shopping list for the Buddy build.

This is the *no-motor* / Pi-only companion build:
camera + mic + speaker + power + amp. All sourced from ecity-iq.com (Iraq).
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


OUT = Path(__file__).resolve().parent.parent / "buddy_electronics_shopping_list.docx"


def set_cell_shading(cell, color_hex: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
    return h


def add_para(doc, text, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    return p


def add_table(doc, headers, rows, col_widths_cm=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for run in hdr[i].paragraphs[0].runs:
            run.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(hdr[i], "1F3A5F")
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
    if col_widths_cm:
        for row in table.rows:
            for i, w in enumerate(col_widths_cm):
                row.cells[i].width = Cm(w)
    return table


# ecity-iq URLs — verified in stock as of build time
ECITY = "https://ecity-iq.com"

ESSENTIALS = [
    ("Pi 5 Power Supply 5.1V 5A 27W USB-C (official spec)", "17,000",
     f"{ECITY}/products/raspberry-pi-5-5-1v-5a-27w-pd-power-supply-usb-c-power-adapter-power-supply"),
    ("Pi 4 Power Supply USB-C (5V ~3A)", "5,500",
     f"{ECITY}/products/power-supply-type-c-charger-for-raspberry-pi-4-model-b"),
    ("Pi Night Vision Fisheye Camera 5MP — 160° FOV (best for face tracking)", "25,000",
     f"{ECITY}/products/raspberry-pi-night-vision-fisheye-camera-5mp-ov5647-160-degree"),
    ("USB Sound Card Adapter (USB → 3.5mm mic + headphone)", "3,000",
     f"{ECITY}/products/external-usb-7-1-sound-card-with-cable-adapter-5hv2-usb-to-3-5mm-jack-3d-ch-sound-audio-headphone-microphone-for-laptop-ps4-012689"),
    ("Mini 2-wire microphone element (plugs into the sound card)", "500",
     f"{ECITY}/products/4x1-5mm-2-wire-mini-microphone-mic-s10-0503275"),
    ("Mono 5W audio amplifier board (drives the speaker)", "1,250",
     f"{ECITY}/search?q=mono+audio+amplifier+5w"),
    ("4Ω 3W speaker 25×35mm (small, fits anywhere)", "1,500",
     f"{ECITY}/collections/speaker"),
    ("USB-C cable with on/off switch (optional — easier power switching)", "2,500",
     f"{ECITY}/products/usb-to-type-c-cable-for-raspberry-pi-4-power-supply-with-on-off-switch"),
]

ALTERNATIVES = [
    ("Pi 5MP standard camera (cheaper, ~53° FOV — face tracking is tight)", "7,500",
     f"{ECITY}/products/raspberry-pi-5mp-1080p-sensor-ov5647-camera-module-for-raspberry-pi-2-3-4-3-010742"),
    ("Pi 5MP + Pi Zero ribbon cable (only for Pi Zero / Pi 5)", "10,000",
     f"{ECITY}/products/raspberry-pi-camera-module-1080p-720p-mini-camera-5mp-with-zero-board-cable"),
]

ACCESSORIES = [
    ("18650 Li-ion 2000mAh cell (for portable battery)", "1,750",
     f"{ECITY}/collections/lithium-battery"),
    ("2S 7.4V 8A BMS protection board (if pairing 2x 18650)", "1,500",
     f"{ECITY}/search?q=2S+BMS"),
    ("TP4056 USB-C charging module", "500",
     f"{ECITY}/products/type-c-usb-5v-1a-tp4056-lithium-battery-charger-module"),
    ("Jumper wire kit (Dupont)", "1,000",
     f"{ECITY}/collections/breadboard-jumper-wires"),
    ("Black aluminum Pi case + fan", "14,000",
     f"{ECITY}/search?q=raspberry+pi+case"),
    ("MicroSD 32GB (if Pi doesn't have one)", "8,000", "any computer shop"),
]


def main():
    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)

    title = doc.add_heading("Buddy Companion Robot — Shopping List", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("No-motor Pi build  ·  All from ecity-iq.com (Erbil-deliverable)")
    r.italic = True
    r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_paragraph()
    add_para(
        doc,
        "Everything below is a direct link to ecity-iq.com (verified in stock). "
        "Pick ONE camera, ONE Pi power supply (matching your Pi model), and EITHER "
        "the USB sound card + mic combo OR a USB headset from any computer shop.",
        italic=True,
    )

    add_heading(doc, "1. Essentials", level=1)
    add_para(
        doc,
        "Buy these to get a working AI companion: power, camera, mic, speaker, amp.",
    )
    add_table(
        doc,
        ["Item", "Price (IQD)", "Link"],
        [[name, price, url] for name, price, url in ESSENTIALS],
        col_widths_cm=[8, 2.5, 6],
    )

    add_para(doc, "")
    add_para(
        doc,
        "Note about the Pi power supply: BUY THE ONE THAT MATCHES YOUR PI."
        " Pi 5 → 5.1V 5A (17,000 IQD). Pi 4 → 5V 3A USB-C (5,500 IQD). Pi 3 / Zero → micro-USB 5V 2.5A.",
        bold=True,
    )

    add_heading(doc, "2. Camera alternatives (cheaper but worse for face tracking)", level=1)
    add_table(
        doc,
        ["Item", "Price (IQD)", "Link"],
        [[name, price, url] for name, price, url in ALTERNATIVES],
        col_widths_cm=[8, 2.5, 6],
    )

    add_heading(doc, "3. Optional add-ons", level=1)
    add_para(
        doc,
        "Skip these for the minimum build. Buy them later if you want portable battery, cooling, or a case.",
    )
    add_table(
        doc,
        ["Item", "Price (IQD)", "Link"],
        [[name, price, url] for name, price, url in ACCESSORIES],
        col_widths_cm=[8, 2.5, 6],
    )

    add_heading(doc, "4. Sourcing in person (Erbil)", level=1)
    add_para(
        doc,
        "If you'd rather go to a shop in Langa Bazaar / Sary Rash, here's what to ask for in Arabic:",
    )
    add_table(
        doc,
        ["What you want", "What to ask for (Arabic)"],
        [
            ["Raspberry Pi power supply 5V", "محول طاقة 5 فولت لراسبيري باي"],
            ["USB conference microphone", "ميكروفون USB"],
            ["USB headset with mic", "هيدسيت USB مع مايك"],
            ["Pi camera module (CSI)", "كاميرا راسبيري باي"],
            ["Small 4Ω 5W speaker", "سپيكر صغير 5 واط"],
            ["Audio amplifier board (small)", "أمبليفاير صوت صغير"],
            ["Dupont jumper wires", "أسلاك جامبر"],
            ["MicroSD card 32GB", "كارت ذاكرة 32 جيجا"],
        ],
        col_widths_cm=[7, 9],
    )

    add_heading(doc, "5. Important checks before ordering", level=1)
    add_para(doc, "✓ Confirm your Raspberry Pi model.", bold=True)
    add_para(
        doc,
        "Pi 4 / 3 / 2 / 1 → uses 15-pin CSI ribbon (default for the cameras above).",
    )
    add_para(
        doc,
        "Pi 5 → uses 22-pin CSI. Buy a 15-to-22-pin adapter ribbon (search 'CSI cable 22 pin' on ecity).",
    )
    add_para(
        doc,
        "Pi Zero / Zero 2 W → uses 22-pin. Buy the 'with ZERO board cable' camera option instead.",
    )
    add_para(doc, "")
    add_para(doc, "✓ Check your webcam first.", bold=True)
    add_para(
        doc,
        "If you already have a USB webcam, plug it into the Pi and run `arecord -l`. "
        "If the webcam has a built-in microphone (most do), you can SKIP buying a mic.",
    )

    doc.save(OUT)
    print(f"Wrote: {OUT}")


if __name__ == "__main__":
    main()
